#!/usr/bin/env python3
"""
Summary Table Generator for Kubenumerate

This module generates a Microsoft Word document containing a summary table
of Kubernetes workloads with security and configuration issues.
"""

from typing import Any, Dict, List, Tuple
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Issue name mapping from technical names to human-readable descriptions
ISSUE_NAME_MAP = {
    "AppArmorNotSet": "AppArmor profile not set",
    "SeccompProfileMissing": "Seccomp profile missing",
    "AutomountServiceAccountTokenTrueAndDefaultSA": "Automounts service account token (default SA)",
    "CapabilityOrSecurityContextMissing": "Missing security context or capabilities",
    "CapabilityAdded": "Additional capabilities granted",
    "CapabilityShouldDropAll": "Should drop all capabilities",
    "DeprecatedAPIUsed": "Uses deprecated API version",
    "LimitsNotSet": "No resource limits set",
    "LimitsCPUNotSet": "No CPU limit set",
    "SensitivePathsMounted": "Sensitive path mounted",
    "RunAsNonRootPSCNilCSCNil": "runAsNonRoot not set",
    "RunAsUserCSCRoot": "Container runs as root (UID 0)",
    "RunAsUserPSCRoot": "Pod runs as root (UID 0)",
    "AllowPrivilegeEscalationNil": "allowPrivilegeEscalation not set",
    "AllowPrivilegeEscalationTrue": "Allows privilege escalation",
    "PrivilegedNil": "Privileged flag not set",
    "PrivilegedTrue": "Running as privileged",
    "ReadOnlyRootFilesystemNil": "Root filesystem not read-only",
    "NamespaceHostPIDTrue": "Uses host PID namespace",
    "NamespaceHostNetworkTrue": "Uses host network namespace",
    "MissingDefaultDenyIngressAndEgressNetworkPolicy": "Missing default deny network policy",
    "AllowAllEgressNetworkPolicyExists": "Allows all egress traffic",
}


def _get_human_readable_issue_name(technical_name: str) -> str:
    """
    Convert technical issue name to human-readable format.

    Args:
        technical_name: Technical issue name from AuditResultName

    Returns:
        Human-readable issue description
    """
    return ISSUE_NAME_MAP.get(technical_name, technical_name)


def aggregate_issues_by_workload(
    findings_df: pd.DataFrame, trivy_results: List[List[Any]], verbosity: int = 1
) -> Dict[Tuple[str, str, str], Dict[str, Any]]:
    """
    Aggregate all issues by workload (namespace, kind, name).

    Excludes ephemeral Pods and groups findings by workload controllers
    like Deployments, DaemonSets, StatefulSets, etc.

    Args:
        findings_df: DataFrame with security/config findings from kubeaudit checks
        trivy_results: List of vulnerable containers from trivy scan
                      Format: [namespace, image_name, pod_name, container_name, crits, crit_cves, highs, high_cves]
        verbosity: Verbosity level for debug output

    Returns:
        Dictionary mapping (namespace, kind, name) to issue summary:
        {
            ("namespace", "Deployment", "app-name"): {
                "issues": ["Human-readable issue 1", "Human-readable issue 2"],
                "issue_count": 2,
                "containers": ["container1", "container2"],
                "vuln_summary": "3 CRITICAL, 5 HIGH CVEs"
            }
        }
    """
    workload_issues: Dict[Tuple[str, str, str], Dict[str, Any]] = {}

    # Define workload controller types (exclude Pod)
    workload_kinds = {
        "Deployment",
        "DaemonSet",
        "StatefulSet",
        "ReplicaSet",
        "Job",
        "CronJob",
        "DeploymentConfig",  # DeploymentConfig for OpenShift
    }

    # Process security/config findings
    if not findings_df.empty:
        for _, row in findings_df.iterrows():
            namespace = str(row.get("ResourceNamespace", "default") or "default")
            kind = str(row.get("ResourceKind", "") or "")
            name = str(row.get("ResourceName", "") or "")
            issue_type = str(row.get("AuditResultName", "") or "")
            container = str(row.get("Container", "") or "")

            # Skip if not a workload controller
            if kind not in workload_kinds:
                continue

            key = (namespace, kind, name)

            # Initialize workload entry if not exists
            if key not in workload_issues:
                workload_issues[key] = {"issues": set(), "containers": set(), "vuln_info": {}}

            # Add human-readable issue
            human_readable_issue = _get_human_readable_issue_name(issue_type)
            workload_issues[key]["issues"].add(human_readable_issue)

            # Track containers
            if container:
                workload_issues[key]["containers"].add(container)

    # Process trivy vulnerability results
    # trivy_results format: [namespace, image_name, pod_name, container_name, crits, crit_cves, highs, high_cves]
    if trivy_results:
        # Group trivy results by namespace and try to map to workloads
        # For now, we'll create a simple mapping by namespace
        namespace_vulns: Dict[str, Dict[str, Any]] = {}

        for vuln_entry in trivy_results:
            if len(vuln_entry) >= 8:
                ns = str(vuln_entry[0]) if vuln_entry[0] is not None else "default"
                pod_name = str(vuln_entry[2]) if vuln_entry[2] is not None else ""
                # Safely convert to int, default to 0 if None or not a number
                try:
                    crits = int(vuln_entry[4]) if vuln_entry[4] is not None else 0
                except (ValueError, TypeError):
                    crits = 0
                try:
                    highs = int(vuln_entry[6]) if vuln_entry[6] is not None else 0
                except (ValueError, TypeError):
                    highs = 0

                if ns not in namespace_vulns:
                    namespace_vulns[ns] = {"crits": 0, "highs": 0, "pods": set()}

                namespace_vulns[ns]["crits"] += crits
                namespace_vulns[ns]["highs"] += highs
                if pod_name:  # Only add if not empty
                    namespace_vulns[ns]["pods"].add(pod_name)

        # Try to match vulnerabilities to workloads
        # This is a heuristic approach - we match by namespace and pod name prefix
        for key in workload_issues.keys():
            namespace, kind, name = key

            if namespace in namespace_vulns:
                # Check if any vulnerable pods match this workload
                # Common pattern: deployment "app" creates pods "app-xxxxx-yyyyy"
                matching_pods = [
                    pod
                    for pod in namespace_vulns[namespace]["pods"]
                    if pod and name and (pod.startswith(name) or name in pod)
                ]

                if matching_pods:
                    crits = namespace_vulns[namespace]["crits"]
                    highs = namespace_vulns[namespace]["highs"]
                    workload_issues[key]["vuln_info"] = {"crits": crits, "highs": highs}

    # Convert sets to lists and add counts
    for key in workload_issues:
        workload_issues[key]["issues"] = sorted(list(workload_issues[key]["issues"]))
        workload_issues[key]["containers"] = sorted(list(workload_issues[key]["containers"]))
        workload_issues[key]["issue_count"] = len(workload_issues[key]["issues"])

        # Add vulnerability summary if present
        vuln_info = workload_issues[key].get("vuln_info", {})
        if vuln_info:
            crits = vuln_info.get("crits", 0)
            highs = vuln_info.get("highs", 0)
            parts = []
            if crits > 0:
                parts.append(f"{crits} CRITICAL")
            if highs > 0:
                parts.append(f"{highs} HIGH")
            if parts:
                vuln_summary = f"Vulnerable images: {', '.join(parts)} CVEs"
                workload_issues[key]["issues"].append(vuln_summary)
                workload_issues[key]["issue_count"] += 1

    # Deduplicate ReplicaSets
    # Remove ReplicaSets if there's a Deployment with same namespace and similar name with identical issues
    # Also remove duplicate ReplicaSets with identical issues (keep only first one)
    deduplicated_issues = {}
    replicaset_signatures: Dict[Tuple[str, Tuple[str, ...]], str] = (
        {}
    )  # Track ReplicaSet issues by namespace to find duplicates

    for key, data in workload_issues.items():
        namespace, kind, name = key

        if kind == "ReplicaSet":
            # Create signature from namespace and sorted issues
            issue_signature = (namespace, tuple(sorted(data["issues"])))

            # Check if there's a Deployment managing this ReplicaSet
            # Deployment name is usually the prefix of ReplicaSet name (before the hash)
            # Example: deployment "nginx" creates replicasets "nginx-1234567890", "nginx-abcdef1234"
            deployment_name = name.rsplit("-", 1)[0] if "-" in name else name
            deployment_key = (namespace, "Deployment", deployment_name)

            # Skip if matching Deployment exists with same issues
            if deployment_key in workload_issues:
                deployment_issues = tuple(sorted(workload_issues[deployment_key]["issues"]))
                replicaset_issues = tuple(sorted(data["issues"]))
                if deployment_issues == replicaset_issues:
                    if verbosity > 1:
                        print(f"[DEBUG] Skipping ReplicaSet {name} - managed by Deployment {deployment_name}")
                    continue

            # Check for duplicate ReplicaSets with identical issues
            if issue_signature in replicaset_signatures:
                # Already have a ReplicaSet with same namespace and issues, skip this one
                if verbosity > 1:
                    existing = replicaset_signatures[issue_signature]
                    print(f"[DEBUG] Skipping duplicate ReplicaSet {name} - same issues as {existing}")
                continue
            else:
                # First ReplicaSet with this signature, keep it
                replicaset_signatures[issue_signature] = name

        # Keep this workload
        deduplicated_issues[key] = data

    if verbosity > 1:
        original_count = len(workload_issues)
        final_count = len(deduplicated_issues)
        if original_count != final_count:
            print(
                f"[DEBUG] Deduplicated {original_count - final_count} ReplicaSets ({final_count} workloads remaining)"
            )

    return deduplicated_issues


def format_issues_list(issues: List[str], max_items: int = 10) -> str:
    """
    Format a list of issues for display in the Word table.

    Args:
        issues: List of human-readable issue descriptions
        max_items: Maximum number of items to show before truncating

    Returns:
        Formatted string with issues
    """
    if not issues:
        return "No issues found"

    if len(issues) <= max_items:
        return "\n".join(f"• {issue}" for issue in issues)
    else:
        displayed = issues[:max_items]
        remaining = len(issues) - max_items
        result = "\n".join(f"• {issue}" for issue in displayed)
        result += f"\n• ... and {remaining} more issues"
        return result


def _set_cell_background(cell: Any, color_rgb: Tuple[int, int, int]) -> None:
    """
    Set background color for a table cell.

    Args:
        cell: docx table cell object
        color_rgb: RGB tuple (r, g, b) with values 0-255
    """
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}")
    cell._element.get_or_add_tcPr().append(shading_elm)


def generate_word_summary_table(
    summary_data: Dict[Tuple[str, str, str], Dict[str, Any]], output_path: str, verbosity: int = 1
) -> None:
    """
    Generate a Microsoft Word document with a summary table of workload issues.

    Args:
        summary_data: Dictionary mapping (namespace, kind, name) to issue data
        output_path: Path where the Word document should be saved
        verbosity: Verbosity level for output messages
    """
    # Create document
    doc = Document()

    # Set document title
    title = doc.add_heading("Kubenumerate Security Issues Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle with info
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        f"This report summarises security and configuration issues found in {len(summary_data)} "
        f"Kubernetes workloads.\n"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    if not summary_data:
        no_issues = doc.add_paragraph()
        no_issues.add_run("✓ No issues detected in workload controllers!").bold = True
        no_issues.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(output_path)
        return

    # Create table with 4 columns: Namespace, Workload, Type, Issues
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"

    # Set column widths
    table.columns[0].width = Inches(1.5)  # Namespace
    table.columns[1].width = Inches(2.0)  # Workload
    table.columns[2].width = Inches(1.2)  # Type
    table.columns[3].width = Inches(3.8)  # Issues

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Namespace"
    header_cells[1].text = "Workload"
    header_cells[2].text = "Type"
    header_cells[3].text = "Issues"

    # Style header row
    for cell in header_cells:
        _set_cell_background(cell, (169, 53, 69))  # Dark red (#A93545)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)

    # Sort workloads by namespace, then kind, then name (ensure all are strings)
    sorted_workloads = sorted(
        summary_data.items(), key=lambda x: (str(x[0][0] or ""), str(x[0][1] or ""), str(x[0][2] or ""))
    )

    # Add data rows
    for idx, ((namespace, kind, name), data) in enumerate(sorted_workloads):
        row_cells = table.add_row().cells

        row_cells[0].text = str(namespace or "")
        row_cells[1].text = str(name or "")
        row_cells[2].text = str(kind or "")
        row_cells[3].text = format_issues_list(data.get("issues", []))

        # Alternate row colors for readability
        if idx % 2 == 0:
            bg_color = (226, 226, 226)  # Light gray (#E2E2E2)
        else:
            bg_color = (255, 255, 255)  # White

        for cell in row_cells:
            _set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Add footer with statistics
    doc.add_paragraph()
    stats = doc.add_paragraph()
    total_issues = sum(data["issue_count"] for data in summary_data.values())
    stats.add_run(f"Summary: Found {total_issues} total issues across {len(summary_data)} workloads.").italic = True

    # Save document
    doc.save(output_path)

    if verbosity > 0:
        print(f"[+] Summary table generated with {len(summary_data)} workloads")


def generate_summary_table(
    findings_df: pd.DataFrame, trivy_results: List[List[Any]], output_path: str, verbosity: int = 1
) -> None:
    """
    Main entry point to generate the summary table.

    Args:
        findings_df: DataFrame with security/config findings
        trivy_results: List of vulnerable containers from trivy
        output_path: Path where the Word document should be saved
        verbosity: Verbosity level for output
    """
    # Aggregate issues by workload
    summary_data = aggregate_issues_by_workload(findings_df, trivy_results, verbosity)

    # Generate Word document
    generate_word_summary_table(summary_data, output_path, verbosity)
